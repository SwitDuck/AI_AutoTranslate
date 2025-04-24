import keras_ocr
import os
import pdf2image
from PyPDF2 import PdfReader
from docx import Document
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline

class ImageToText:
    def __init__(self, destination):
        self.destination = destination
        self.pipeline = keras_ocr.pipeline.Pipeline()
        
    def PdfToImage(self):
        pdf_path = os.path.join(self.destination, 'book.pdf')
        os.makedirs(os.path.join(self.destination, 'images'), exist_ok=True)

        pages = pdf2image.convert_from_path(
            pdf_path,
            output_folder=os.path.join(self.destination, 'images'),
            dpi=200,
            poppler_path='A:\\poppler-24.08.0\\Library\\bin',
            fmt='jpeg',
            jpegopt={"quality": 85},
        )

        text_file_path = os.path.join(self.destination, 'TextFile.txt')
        with open(text_file_path, 'w', encoding='utf-8') as f:
            for idx, page in enumerate(pages):
                img_path = os.path.join(self.destination, 'images', f'page{idx}.jpg')
                page.save(img_path, 'JPEG')
                prediction_groups = self.pipeline.recognize([img_path])
                text = ' '.join([word for word, box in prediction_groups[0]])
                f.write(text + '\n')
                os.remove(img_path)
    def PdfTextAnalize(self):
        reader = PdfReader(self.destination+'\\book.pdf')
        pages = reader.pages
        doc = Document()
        doc.add_heading('Обработанный текст', level=1)
        doc.add_paragraph('Текст из pdf файла')
        def visitor_body(text, cm, tm, fontDict, fontSize):
            y = tm[5]
            if y > 50 and y < 720:
                return text
        for page in pages:
            text = page.extract_text(visitor_text=visitor_body)
            doc.add_paragraph(text)
        doc.save("TextFromBook.docx")

class AiAPI:
    def __init__(self, apikey):
        self.apikey = apikey
    '''
    def AiPrompting(self):
        from openai import OpenAI
        client = OpenAI()
        response = client.responses.create(
            model="gpt-4.1",
            input=[
                {
                    "role": "user",
                    "content": "You will translate from en to ru text i prompt you, text might include short techinque word and words not legible because of tokens amound, skip them or try to anticipate, here text"+
                }
            ]
        )
        print(response.output_text)
    
    def AiReqHandler(self):
        ...
    '''