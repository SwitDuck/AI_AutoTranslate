import asyncio
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
from docx import Document
import os
import re

class AsyncDocxTranslator:
    def __init__(self):
        self.model_name = "facebook/nllb-200-distilled-600M"
        self.tokenizer = AutoTokenizer.from_pretrained(self.model_name)
        self.model = AutoModelForSeq2SeqLM.from_pretrained(self.model_name)
        self.translator = pipeline(
            "translation",
            model=self.model,
            tokenizer=self.tokenizer,
            src_lang="eng_Latn",
            tgt_lang="rus_Cyrl",
            device="cpu"
        )
        self.max_length = 512
        self.batch_size = 10
        self.temp_file = "temp_translated.docx"

    def split_long_text(self, text):
        """Разбивает длинный текст на части по предложениям"""
        sentences = re.split(r'(?<=[.!?]) +', text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) < self.max_length * 0.8:
                current_chunk += " " + sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
        
        if current_chunk:
            chunks.append(current_chunk.strip())
            
        return chunks

    async def translate_text(self, text):
        """Асинхронный перевод текста"""
        if len(text) > self.max_length * 0.9:
            chunks = self.split_long_text(text)
            translated_chunks = []
            
            for chunk in chunks:
                loop = asyncio.get_event_loop()
                result = await loop.run_in_executor(
                    None, 
                    lambda: self.translator(chunk, max_length=self.max_length)[0]["translation_text"]
                )
                translated_chunks.append(result)
                await asyncio.sleep(0.1)
            
            return " ".join(translated_chunks)
        else:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(
                None, 
                lambda: self.translator(text, max_length=self.max_length)[0]["translation_text"]
            )

    async def process_paragraphs(self, paragraphs, start_page=0, end_page=None, existing_doc=None):
        """Обработка абзацев с добавлением к существующему документу"""
        translated_doc = existing_doc or Document()
        end_page = end_page if end_page is not None else len(paragraphs)
        
        for i, paragraph in enumerate(paragraphs[start_page:end_page], start=start_page):
            if paragraph.text.strip():
                print(f"Обработка абзаца {i + 1}/{len(paragraphs)}...")
                translated_text = await self.translate_text(paragraph.text)
                translated_doc.add_paragraph(translated_text)
                
                if (i + 1) % self.batch_size == 0:
                    print(f"Промежуточное сохранение...")
                    translated_doc.save(self.temp_file)
                    await asyncio.sleep(1)

        return translated_doc

    async def translate_docx(self, input_path, output_path, start_page=0, end_page=None):
        """Перевод с добавлением к существующему файлу"""
        # Загружаем существующий перевод или создаем новый
        if os.path.exists(output_path):
            print(f"Обнаружен существующий файл перевода. Добавляем к нему...")
            existing_doc = Document(output_path)
        else:
            print("Создаем новый файл перевода...")
            existing_doc = Document()

        # Загружаем исходный документ
        source_doc = Document(input_path)

        # Обрабатываем выбранные страницы
        translated_doc = await self.process_paragraphs(
            source_doc.paragraphs,
            start_page=start_page,
            end_page=end_page,
            existing_doc=existing_doc
        )
        
        # Сохраняем результат
        translated_doc.save(output_path)
        if os.path.exists(self.temp_file):
            os.remove(self.temp_file)
        
        print(f"Перевод завершён! Результат сохранён в {output_path}")